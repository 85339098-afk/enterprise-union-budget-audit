#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工会预决算审核意见文档 — 审核业务内容生成模块

本模块专注审核业务内容的生成（表格、问题清单、结论等），
不包含页面设置、字体、页码等排版逻辑——这些由"公文排版"skill 的 API 处理。
"""

from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import platform
FONT_MAP = {
    "仿宋_GB2312": "仿宋",
    "楷体_GB2312": "楷体",
}
def resolve_font_name(fn):
    if platform.system() == "Darwin":
        return FONT_MAP.get(fn, fn)
    return fn





def add_audit_heading(doc, union_name, year_start=2025, year_end=2026):
    """
    添加审核意见专用标题占位函数。

    实际排版由“公文排版”skill 的 add_title() 完成（加载 gov-doc-format 后调用）。
    本函数仅作为语义化占位，AI编排时直接调用 gov-doc-format 的 add_title()。

    Args:
        doc: Document 对象
        union_name: 工会全称
        year_start: 决算年度（默认2025）
        year_end: 预算年度（默认2026）
    """
    pass


def _set_cell_font(cell, font_name="仿宋_GB2312", font_size=14, bold=False):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(24)
        for run in paragraph.runs:
            run.font.name = resolve_font_name(font_name)
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(0, 0, 0)


def _shade_cells(row, color_hex="D9E2F3"):
    for cell in row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color_hex)
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)


def add_review_table(doc, items):
    """添加整改复核对照表"""
    if not items:
        return
    headers = ["序号", "原审核意见", "整改情况", "复核结论"]
    table = doc.add_table(rows=1 + len(items), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_row = table.rows[0]
    _shade_cells(header_row, "D9E2F3")
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        cell.paragraphs[0].add_run(header)
        _set_cell_font(cell, "黑体", 14)
    for idx, item in enumerate(items):
        row = table.rows[idx + 1]
        values = [
            str(item.get("no", idx + 1)),
            item.get("original", ""),
            item.get("status", ""),
            item.get("conclusion", ""),
        ]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ""
            cell.paragraphs[0].add_run(val)
            _set_cell_font(cell, "仿宋_GB2312", 12)
    doc.add_paragraph()


def add_new_issue(doc, issue):
    """添加单个新问题的结构化描述"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("问题" + str(issue.get("no", "")) + "：" + issue.get("title", ""))
    run.font.name = resolve_font_name("黑体")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), resolve_font_name("黑体"))
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)
    labels = [
        ("reports", "涉及报表"),
        ("location", "具体位置"),
        ("description", "问题描述"),
        ("basis", "审核依据"),
        ("suggestion", "整改建议"),
    ]
    for key, label in labels:
        val = issue.get(key)
        if val:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = Pt(28)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r1 = p.add_run("【" + label + "】")
            r1.font.name = resolve_font_name("楷体_GB2312")
            r1._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), resolve_font_name("楷体_GB2312"))
            r1.font.size = Pt(16)
            r1.font.bold = True
            r1.font.color.rgb = RGBColor(0, 0, 0)
            r2 = p.add_run(val)
            r2.font.name = resolve_font_name("仿宋_GB2312")
            r2._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), resolve_font_name("仿宋_GB2312"))
            r2.font.size = Pt(16)
            r2.font.color.rgb = RGBColor(0, 0, 0)


def add_issue_category(doc, category):
    """添加问题类别分组"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(category.get("name", ""))
    run.font.name = resolve_font_name("楷体_GB2312")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), resolve_font_name("楷体_GB2312"))
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    for issue in category.get("issues", []):
        add_new_issue(doc, issue)


def add_conclusion_section(doc, conclusions):
    """添加审核结论段落"""
    if isinstance(conclusions, str):
        conclusions = [conclusions]
    for c in conclusions:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(c)
        run.font.name = resolve_font_name("仿宋_GB2312")
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), resolve_font_name("仿宋_GB2312"))
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_requirement_section(doc, requirements):
    """添加整改要求段落"""
    for req in requirements:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        text = str(req.get("no", "")) + ". " + req.get("text", "")
        run = p.add_run(text)
        run.font.name = resolve_font_name("仿宋_GB2312")
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), resolve_font_name("仿宋_GB2312"))
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_special_note_section(doc, pending_items):
    """添加'需核实事项'段落"""
    for idx, item in enumerate(pending_items):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        text = str(idx + 1) + ". " + item.get("item", "") + "（原因：" + item.get("reason", "需补充说明") + "）"
        run = p.add_run(text)
        run.font.name = resolve_font_name("仿宋_GB2312")
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), resolve_font_name("仿宋_GB2312"))
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)


if __name__ == "__main__":
    from docx import Document
    d = Document()
    add_review_table(d, [{"no": 1, "original": "test", "status": "ok", "conclusion": "done"}])
    add_issue_category(d, {"name": "Cat", "issues": [{"no": 1, "title": "T", "reports": "R", "location": "L", "description": "D", "basis": "B", "suggestion": "S"}]})
    add_conclusion_section(d, ["Done"])
    add_requirement_section(d, [{"no": 1, "text": "Fix it"}])
    add_special_note_section(d, [{"item": "Check X", "reason": "unclear"}])
    print("Self-test passed.")
