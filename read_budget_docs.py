#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工会预算审核文档读取脚本
读取 docx、xlsx、pdf 格式的预决算报表和审核意见。
带依赖回退机制：markitdown 不可用时自动降级到 python-docx 原生读取。
"""

import sys
import os
import traceback
from pathlib import Path


# ============================================================
# DOCX 读取（支持 markitdown / python-docx 双模式）
# ============================================================

def read_docx_markitdown(file_path):
    """使用 markitdown 读取（首选）"""
    from markitdown import MarkItDown
    md = MarkItDown()
    result = md.convert(file_path)
    return result.text_content


def read_docx_fallback(file_path):
    """回退：使用 python-docx 原生读取"""
    from docx import Document
    doc = Document(file_path)
    output = []
    for para in doc.paragraphs:
        if para.text.strip():
            output.append(para.text)
    # 读取表格
    for table in doc.tables:
        output.append("\n[表格开始]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            output.append(" | ".join(cells))
        output.append("[表格结束]\n")
    return "\n".join(output)


def read_docx(file_path):
    """读取Word文档，含依赖回退"""
    try:
        return read_docx_markitdown(file_path)
    except ImportError:
        print("[INFO] markitdown 不可用，降级为 python-docx 原生读取", file=sys.stderr)
        return read_docx_fallback(file_path)
    except Exception as e:
        # markitdown 可能成功但报错，尝试回退
        try:
            print(f"[WARN] markitdown 读取异常: {e}，尝试降级", file=sys.stderr)
            return read_docx_fallback(file_path)
        except Exception as e2:
            raise RuntimeError(f"DOCX 读取失败（markitdown: {e}，回退: {e2}）") from e


# ============================================================
# XLSX 读取
# ============================================================

def read_xlsx(file_path):
    """读取Excel报表（所有Sheet）"""
    try:
        import pandas as pd
    except ImportError:
        raise RuntimeError("缺少 pandas/openpyxl 依赖，请执行: pip install pandas openpyxl")

    all_sheets = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
    output = []
    for sheet_name, df in all_sheets.items():
        output.append(f"\n{'=' * 60}")
        output.append(f"Sheet: {sheet_name}")
        output.append(f"{'=' * 60}\n")
        # 将 DataFrame 转为文本，空值显示为空字符串
        output.append(df.fillna('').to_string(index=False))
        output.append("\n")
    return "\n".join(output)


# ============================================================
# PDF 读取
# ============================================================

def read_pdf(file_path):
    """读取PDF文档"""
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("缺少 pdfplumber 依赖，请执行: pip install pdfplumber")

    output = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            output.append(f"\n{'=' * 60}")
            output.append(f"Page {i + 1}")
            output.append(f"{'=' * 60}\n")
            text = page.extract_text()
            if text:
                output.append(text)
            else:
                output.append("（本页无可提取文本，可能为扫描图片）")
            output.append("\n")
    return "\n".join(output)


# ============================================================
# 主入口
# ============================================================

SUPPORTED_EXTENSIONS = {
    '.docx': 'Word 文档',
    '.xlsx': 'Excel 报表',
    '.pdf': 'PDF 文档',
}


def main():
    if len(sys.argv) < 2:
        print("用法: python read_budget_docs.py <文件路径>", file=sys.stderr)
        print("支持的格式: .docx .xlsx .pdf", file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]

    if not os.path.exists(file_path):
        print(f"错误: 文件不存在: {file_path}", file=sys.stderr)
        sys.exit(1)

    ext = Path(file_path).suffix.lower()

    readers = {
        '.docx': read_docx,
        '.xlsx': read_xlsx,
        '.pdf': read_pdf,
    }

    reader = readers.get(ext)
    if reader is None:
        print(f"错误: 不支持的文件格式: {ext}", file=sys.stderr)
        print(f"支持的格式: {', '.join(SUPPORTED_EXTENSIONS.keys())}", file=sys.stderr)
        sys.exit(1)

    try:
        content = reader(file_path)
        print(content)
    except RuntimeError as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"读取失败: {e}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
